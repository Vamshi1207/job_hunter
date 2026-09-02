"""Build Word and Pages copies of a tailored HTML resume so the PDF is not the only edit path."""

from __future__ import annotations

import json
import logging
import os
import subprocess
import sys
import threading
import time
import urllib.error
import urllib.request
import zipfile
from pathlib import Path

from bs4 import BeautifulSoup, Tag

log = logging.getLogger(__name__)

SKIP_TAGS = {"script", "style", "head", "meta", "link", "title"}
PAGES_HELPER_DEFAULT = "http://host.docker.internal:8765"
_HELPER_OK: bool | None = None
_HELPER_OK_UNTIL = 0.0
_PAGES_APP_LOCK = threading.Lock()


def write_editable_exports(html_path: Path, pdf_path: Path | None = None) -> dict[str, Path]:
    """Write .docx and a native .pages file next to the HTML resume. Returns paths that exist."""
    html_path = Path(html_path)
    if not html_path.is_file():
        return {}
    stem_path = html_path.with_suffix("")
    docx_path = Path(str(stem_path) + ".docx")
    pages_path = Path(str(stem_path) + ".pages")
    written: dict[str, Path] = {"html": html_path}
    try:
        html_to_docx(html_path, docx_path)
        written["docx"] = docx_path
    except Exception as exc:
        log.warning("Word export failed for %s: %s", html_path.name, exc)
    try:
        html_to_pages(html_path, pages_path, pdf_path=pdf_path, docx_path=docx_path)
    except Exception as exc:
        log.warning("Pages export failed for %s: %s", html_path.name, exc)
    if is_native_pages(pages_path):
        written["pages"] = pages_path
    return written


def ensure_package_exports(folder: Path) -> dict[str, str | None]:
    """Create missing Word/Pages files for an application folder. Safe to call on list."""
    folder = Path(folder)
    html = next(iter(sorted(folder.glob("*_CV.html"))), None)
    names = {
        "html_name": html.name if html else None,
        "docx_name": None,
        "pages_name": None,
    }
    docx = html.with_suffix(".docx") if html else next(iter(sorted(folder.glob("*_CV.docx"))), None)
    pages = Path(str(html.with_suffix("")) + ".pages") if html else next(iter(sorted(folder.glob("*_CV.pages"))), None)
    if is_fake_pages(pages):
        try:
            Path(pages).unlink()
        except OSError:
            pass
    if html and (not docx or not Path(docx).is_file()):
        try:
            html_to_docx(html, Path(docx))
        except Exception as exc:
            log.warning("Word export failed for %s: %s", html.name, exc)
    if html and docx and Path(docx).is_file() and not is_native_pages(pages):
        try:
            html_to_pages(html, Path(pages), docx_path=Path(docx))
        except Exception as exc:
            log.warning("Pages export failed for %s: %s", html.name, exc)
    if docx and Path(docx).is_file():
        names["docx_name"] = Path(docx).name
    if is_native_pages(pages):
        names["pages_name"] = Path(pages).name
    return names


def html_to_docx(html_path: Path, docx_path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt, RGBColor

    soup = BeautifulSoup(html_path.read_text(encoding="utf-8"), "html.parser")
    doc = Document()
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    navy = RGBColor(0x1E, 0x3A, 0x5F)
    ink = RGBColor(0x1A, 0x1A, 0x1A)
    muted = RGBColor(0x5C, 0x5C, 0x5C)

    def _set_run(run, *, size=11, bold=False, italic=False, color=ink, name="Calibri"):
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.name = name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)

    def _para(text, *, size=11, bold=False, italic=False, color=ink, space_after=4, space_before=0, align=None):
        text = (text or "").strip()
        if not text:
            return None
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        if align:
            p.alignment = align
        run = p.add_run(text)
        _set_run(run, size=size, bold=bold, italic=italic, color=color)
        return p

    body = soup.body or soup
    for el in body.find_all(True):
        if el.name in SKIP_TAGS:
            el.decompose()

    for el in body.children:
        if not isinstance(el, Tag):
            continue
        classes = el.get("class") or []
        if "header" in classes:
            _para(_text(el.select_one(".header-name")), size=22, bold=True, color=navy, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
            _para(_text(el.select_one(".header-title")), size=11, italic=True, space_after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
            _para(_contact_text(el.select_one(".header-contact")), size=10, color=muted, space_after=10, align=WD_ALIGN_PARAGRAPH.CENTER)
            continue
        if "summary" in classes:
            _para(_text(el), size=11, space_after=10)
            continue
        if "section-heading" in classes:
            _para(_text(el), size=11, bold=True, color=navy, space_before=10, space_after=6)
            continue
        if "job-block" in classes or "school-block" in classes or "project-block" in classes:
            _add_job_block(el, _para, doc, _set_run, navy, ink, muted)
            continue
        if "skills-list" in classes:
            for row in _skill_rows(el):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                label = p.add_run(row[0] + ": ")
                _set_run(label, size=10, bold=True, color=navy)
                items = p.add_run(row[1])
                _set_run(items, size=10, color=ink)
            continue
        if "keep-skills" in classes or "keep-education" in classes:
            heading = el.select_one(".section-heading")
            if heading:
                _para(_text(heading), size=11, bold=True, color=navy, space_before=10, space_after=6)
            skills = el.select_one(".skills-list")
            if skills:
                for row in _skill_rows(skills):
                    p = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    label = p.add_run(row[0] + ": ")
                    _set_run(label, size=10, bold=True, color=navy)
                    items = p.add_run(row[1])
                    _set_run(items, size=10, color=ink)
            for block in el.select(".school-block, .job-block, .project-block"):
                _add_job_block(block, _para, doc, _set_run, navy, ink, muted)
            continue

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))


def html_to_pages(
    html_path: Path,
    pages_path: Path,
    pdf_path: Path | None = None,
    docx_path: Path | None = None,
) -> None:
    """Write a native Pages document by asking Pages.app to save the Word copy."""
    del pdf_path  # preview lives inside the Word file; Pages.app builds its own
    pages_path = Path(pages_path)
    if is_native_pages(pages_path):
        return
    if is_fake_pages(pages_path):
        pages_path.unlink()
    docx_path = Path(docx_path) if docx_path else Path(html_path).with_suffix(".docx")
    if not docx_path.is_file():
        html_to_docx(Path(html_path), docx_path)
    docx_to_pages(docx_path, pages_path)


def is_native_pages(path: Path | None) -> bool:
    """True when the file is a real Pages document (IWA), not a renamed zip of HTML."""
    if path is None:
        return False
    path = Path(path)
    if path.is_dir():
        return (path / "Index" / "Document.iwa").is_file()
    if not path.is_file():
        return False
    try:
        with zipfile.ZipFile(path) as zf:
            names = zf.namelist()
    except (OSError, zipfile.BadZipFile):
        return False
    return any(name == "Index/Document.iwa" or name.endswith("/Document.iwa") for name in names)


def is_fake_pages(path: Path | None) -> bool:
    """True when a .pages zip exists but Pages.app cannot open it."""
    if path is None:
        return False
    path = Path(path)
    if not path.is_file() or is_native_pages(path):
        return False
    try:
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
    except (OSError, zipfile.BadZipFile):
        return True
    return "Index/Document.iwa" not in names and not any(n.endswith("/Document.iwa") for n in names)


def docx_to_pages(docx_path: Path, pages_path: Path) -> None:
    """Convert Word to native Pages via Pages.app (Mac) or the host helper (Docker)."""
    if os.environ.get("JOB_SEARCH_SKIP_PAGES_APP", "").lower() in {"1", "true", "yes"}:
        raise RuntimeError("Pages export skipped")
    docx_path = Path(docx_path)
    pages_path = Path(pages_path)
    if not docx_path.is_file():
        raise FileNotFoundError(docx_path)
    if _pages_app_enabled():
        docx_to_pages_via_app(docx_path, pages_path)
        return
    helper = (os.environ.get("JOB_SEARCH_PAGES_HELPER") or "").strip()
    if not helper and _in_docker():
        helper = PAGES_HELPER_DEFAULT
    if helper and _helper_available(helper):
        _docx_to_pages_via_helper(helper, docx_path, pages_path)
        return
    raise RuntimeError("Pages.app is not available in this environment")


def docx_to_pages_via_app(docx_path: Path, pages_path: Path) -> None:
    """Open a Word file in Pages.app and save it as a native .pages document."""
    docx_path = Path(docx_path).resolve()
    pages_path = Path(pages_path).resolve()
    pages_path.parent.mkdir(parents=True, exist_ok=True)
    if pages_path.exists():
        if pages_path.is_dir():
            raise RuntimeError(f"Refusing to replace directory {pages_path}")
        pages_path.unlink()
    src = _applescript_posix(docx_path)
    dest = _applescript_posix(pages_path)
    script = (
        f'set src to POSIX file "{src}"\n'
        f'set dest to POSIX file "{dest}"\n'
        'tell application "Pages"\n'
        "  set theDoc to open src\n"
        "  delay 1\n"
        "  try\n"
        "    close theDoc saving in dest\n"
        "  on error errMsg\n"
        "    try\n"
        "      close theDoc saving no\n"
        "    end try\n"
        "    error errMsg\n"
        "  end try\n"
        "end tell\n"
    )
    with _PAGES_APP_LOCK:
        last_err = ""
        for attempt in range(3):
            _ensure_pages_app()
            if attempt:
                time.sleep(0.8 * attempt)
            try:
                result = subprocess.run(
                    ["osascript", "-e", script],
                    capture_output=True,
                    text=True,
                    timeout=25,
                )
            except subprocess.TimeoutExpired:
                last_err = "Pages AppleScript timed out"
                continue
            if result.returncode == 0:
                break
            last_err = (result.stderr or result.stdout or "").strip()
            if not _pages_retryable(last_err):
                raise RuntimeError(last_err or "Pages AppleScript failed")
        else:
            raise RuntimeError(last_err or "Pages.app did not start")
    if not is_native_pages(pages_path):
        raise RuntimeError(f"Pages.app did not write a native document at {pages_path}")


def _ensure_pages_app() -> None:
    """Start Pages.app without stealing focus. AppleScript `tell` will not launch it from a helper."""
    subprocess.run(
        ["open", "-g", "-a", "Pages"],
        check=False,
        timeout=20,
        capture_output=True,
    )
    deadline = time.monotonic() + 12
    while time.monotonic() < deadline:
        try:
            result = subprocess.run(
                ["osascript", "-e", 'tell application "Pages" to get version'],
                capture_output=True,
                text=True,
                timeout=4,
            )
        except subprocess.TimeoutExpired:
            time.sleep(0.5)
            continue
        if result.returncode == 0:
            return
        if not _pages_retryable((result.stderr or result.stdout or "").strip()):
            return
        time.sleep(0.4)


def _pages_retryable(message: str) -> bool:
    text = message.lower()
    return (
        "-600" in message
        or "-609" in message
        or "-2700" in message
        or "isn't running" in text
        or "isn’t running" in text
        or "not running" in text
        or "connection is invalid" in text
    )


def _pages_app_enabled() -> bool:
    if os.environ.get("JOB_SEARCH_SKIP_PAGES_APP", "").lower() in {"1", "true", "yes"}:
        return False
    if _in_docker() or sys.platform != "darwin":
        return False
    return Path("/Applications/Pages.app").is_dir()


def _in_docker() -> bool:
    return os.environ.get("IN_DOCKER") == "1" or Path("/.dockerenv").exists()


def _applescript_posix(path: Path) -> str:
    return str(path).replace("\\", "\\\\").replace('"', '\\"')


def _host_path(path: Path) -> str:
    host_root = (os.environ.get("HOST_PROJECT_ROOT") or "").strip()
    job_root = (os.environ.get("JOB_SEARCH_ROOT") or "").strip()
    resolved = path.resolve()
    if host_root and job_root:
        try:
            rel = resolved.relative_to(Path(job_root).resolve())
            return str(Path(host_root) / rel)
        except ValueError:
            pass
    return str(resolved)


def _helper_available(helper: str) -> bool:
    global _HELPER_OK, _HELPER_OK_UNTIL
    now = time.monotonic()
    if _HELPER_OK is not None and now < _HELPER_OK_UNTIL:
        return _HELPER_OK
    try:
        urllib.request.urlopen(helper.rstrip("/") + "/health", timeout=0.4)
        _HELPER_OK = True
        _HELPER_OK_UNTIL = now + 60
    except Exception:
        _HELPER_OK = False
        _HELPER_OK_UNTIL = now + 15
    return _HELPER_OK


def _docx_to_pages_via_helper(helper: str, docx_path: Path, pages_path: Path) -> None:
    url = helper.rstrip("/") + "/pages"
    payload = json.dumps(
        {"docx": _host_path(docx_path), "pages": _host_path(pages_path)}
    ).encode()
    req = urllib.request.Request(
        url, data=payload, method="POST", headers={"Content-Type": "application/json"}
    )
    try:
        with urllib.request.urlopen(req, timeout=90) as resp:
            if resp.status >= 400:
                raise RuntimeError(f"Pages helper HTTP {resp.status}")
    except urllib.error.HTTPError as exc:
        detail = ""
        try:
            detail = exc.read().decode("utf-8", errors="replace").strip()[:500]
        except Exception:
            detail = str(exc.reason or exc)
        raise RuntimeError(f"Pages helper HTTP {exc.code} at {url}: {detail or exc.reason}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"Pages helper unreachable at {url}: {exc}") from exc
    if not is_native_pages(pages_path):
        raise RuntimeError("Pages helper did not write a native document")


def _text(el) -> str:
    if el is None:
        return ""
    return " ".join(el.get_text(" ", strip=True).split())


def _contact_text(el) -> str:
    if el is None:
        return ""
    parts = []
    for child in el.children:
        if isinstance(child, Tag):
            if "dot" in (child.get("class") or []):
                parts.append("·")
            else:
                t = _text(child)
                if t:
                    parts.append(t)
        else:
            t = str(child).strip()
            if t:
                parts.append(t)
    return " ".join(parts) if parts else _text(el)


def _skill_rows(el: Tag) -> list[tuple[str, str]]:
    labels = el.select(".skill-label")
    items = el.select(".skill-items")
    rows = []
    for label, item in zip(labels, items):
        rows.append((_text(label).rstrip(":"), _text(item)))
    return rows


def _add_job_block(el: Tag, _para, doc, _set_run, navy, ink, muted) -> None:
    from docx.shared import Pt

    headers = el.select(".job-header")
    for header in headers:
        left = header.select_one(".company-name, .job-title, .job-main")
        right = header.select_one(".job-place, .job-dates, .proj-link")
        left_text = _text(left) if left else _text(header)
        right_text = _text(right) if right else ""
        if not left_text and not right_text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        is_company = bool(header.select_one(".company-name"))
        run = p.add_run(left_text)
        _set_run(run, size=11, bold=is_company, italic=not is_company and bool(header.select_one(".job-title")), color=navy if is_company else ink)
        if right_text:
            gap = p.add_run("    " + right_text)
            _set_run(gap, size=9, color=muted)
    stack = el.select_one(".project-stack, .edu-focus")
    if stack:
        _para(_text(stack), size=9, italic=True, color=muted, space_after=2)
    for li in el.select("li"):
        point = li.select_one(".point") or li
        text = _text(point)
        if not text:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("• " + text)
        _set_run(run, size=10, color=ink)
